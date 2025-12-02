# gradio_app.py
import gradio as gr, os, shutil, tempfile
from datetime import datetime
from typing import Tuple, List
from rag_chatbot import (ingest_documents, retrieve_relevant, build_prompt,
                         generate_essay, add_documents_to_store, export_store, import_store)

# For file export
from docx import Document
from fpdf import FPDF
from fpdf.enums import XPos, YPos
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ----------  GLOBAL VARIABLES  ----------
PendingEntry = Tuple[str, str]          # (real_name, temp_path)
pending_files = gr.State()
pending_files.value: List[PendingEntry] = []   # list of absolute paths not yet indexed 

store = None
stored_files = []          # list of (filename, filepath) tuples already in store
last_essay = None
last_refs = None

# ----------  DEFAULT UI VALUES  ----------
pending_value = "No files in queue"

def keep_upload(gradio_path: str) -> str:
    """Keep a permanent copy of the file Gradio already saved for us."""
    permanent = tempfile.NamedTemporaryFile(delete=False,
                                          suffix=os.path.splitext(gradio_path)[1])
    shutil.copy2(gradio_path, permanent.name)
    return permanent.name

def save_uploaded_files(files):
    """Save uploaded files and return their paths."""
    if not files:
        return []
    
    paths = []
    for f in files:
        if isinstance(f, str):
            # Make a permanent copy
            suffix = os.path.splitext(f)[1]
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            with open(f, 'rb') as src:
                tmp.write(src.read())
            tmp.close()
            paths.append((os.path.basename(f), tmp.name))
        else:
            suffix = "." + f.name.split(".")[-1]
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            tmp.write(f.read())
            tmp.close()
            paths.append((f.name, tmp.name))
    
    return paths

# ----------  NEW SMALL HELPER  ----------
def base_name(entries: List[PendingEntry]) -> List[str]:
    """Extract real names for display."""
    return [name for name, _ in entries]

# ----------  MODIFIED  "process"  LOGIC  ----------
def process_files() -> tuple:
    """Ingest everything that is currently queued."""
    global store, stored_files          # stored_files is still List[(name,path)]

    if not pending_files.value:
        return format_file_list(stored_files), "❗ No pending files to add.", gr.update(value=pending_value)

    paths = [tmp for _, tmp in pending_files.value]   # ingest needs paths only
    names = [name for name, _ in pending_files.value] # keep names for UI

    if store is None:                     # first batch
        store, _ = ingest_documents(paths)
        stored_files.extend([(n, p) for n, p in zip(names, paths)])
        msg = f"✅ Indexed {len(store.metadatas)} chunks from {len(paths)} file(s)."
    else:                                 # incremental
        num_chunks, _ = add_documents_to_store(store, paths)
        stored_files.extend([(n, p) for n, p in zip(names, paths)])
        msg = f"✅ Added {num_chunks} chunks. Total: {len(store.metadatas)}."

    pending_files.value = []              # empty queue
    return format_file_list(stored_files), msg, gr.update(value=pending_value)

# ----------  NEW DROP HANDLER  ----------
def on_drop_more(files) -> tuple:
    """Append newly dropped files to the pending list + clear the box."""
    if not files:
        return pending_files, gr.update(value="\n".join(base_name(pending_files.value))), gr.update(value=None)

    new_entries = [(os.path.basename(f), keep_upload(f)) for f in files]
    pending_files.value.extend(new_entries)

    return pending_files, \
           gr.update(value="\n".join(base_name(pending_files.value))), \
           gr.update()

def format_file_list(entries: List[tuple[str, str]]) -> str:
    """Pretty list of real names."""
    if not entries:
        return "No files"
    return "\n".join(f"{i}. {name}" for i, (name, _) in enumerate(entries))

def remove_file(index_str):
    """Remove a specific file from the store."""
    global store, stored_files
    
    if not index_str or not stored_files:
        return format_file_list(stored_files), "❗ Invalid index or no files to remove."
    
    try:
        idx = int(index_str)
        if idx < 0 or idx >= len(stored_files):
            return format_file_list(stored_files), "❗ Index out of range."
        
        # Remove file
        name, path = stored_files.pop(idx)
        try:
            os.remove(path)
        except:
            pass
        
        # Rebuild store without this file
        if stored_files:
            remaining_paths = [path for _, path in stored_files]
            store, _ = ingest_documents(remaining_paths)
            message = f"✅ Removed '{name}'. Store rebuilt with {len(stored_files)} file(s)."
        else:
            store = None
            message = f"✅ Removed '{name}'. Store is now empty."
        
        return format_file_list(stored_files), message
        
    except ValueError:
        return format_file_list(stored_files), "❗ Please enter a valid number."

def clear_all():
    """Clear all files from the store."""
    global store, stored_files
    
    # Clean up temp files
    for _, path in stored_files:
        try:
            os.remove(path)
        except:
            pass
    
    stored_files = []
    store = None
    
    return "No files in store", "🗑️ All files cleared from store."

# ----------  SAVE / LOAD  ----------
def save_store_file() -> str:
    """Return *path* of the saved store (gzipped json)."""
    if store is None:
        gr.Warning("No store to save")
        return ""                # empty path → button disabled
    
    blob = export_store(store)
    tmp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".vsto.json").name
    with open(tmp_path, "wb") as f:
        f.write(blob)
    return tmp_path              # Gradio offers this file for download

def load_store_file(file) -> tuple:
    """Called by gr.File – uploads a *.vsto.json store."""
    global store, stored_files
    if file is None:
        return format_file_list(stored_files), "❗ No file selected."

    try:
        # Gradio gives either str (path) or bytes; handle both
        if isinstance(file, str):
            with open(file, "rb") as f:
                blob = f.read()
        else:
            blob = file.read() if hasattr(file, "read") else file

        store = import_store(blob)
        # rebuild UI list from metadata (we only have basename in meta)
        stored_files = [(meta["source"], "loaded") for meta in store.metadatas]
        # (optional) give unique dummy paths if you need them later
        return format_file_list(stored_files), f"✅ Loaded store with {len(store.metadatas)} chunks."
    except ValueError as e:
        return format_file_list(stored_files), f"❌ {e}"

def generate_paper(query):
    global store, last_essay, last_refs
    
    if store is None:
        return "❗ Please upload and process your notes first.", gr.update(visible=False)
    
    try:
        context, hits = retrieve_relevant(store, query, k=6)
        prompt = build_prompt(context, query)
        essay = generate_essay(prompt)
        
        if not essay or essay.strip() == "":
            return "❗ Generated essay is empty. Please try again.", gr.update(visible=False)
        
        refs = "\n".join([f"- {h[1]['source']} (chunk {h[1]['chunk']}, score={h[0]:.3f})" for h in hits])
        
        # Store for export
        last_essay = essay
        last_refs = refs
        
        #result = f"### 📄 Generated Academic Paper\n\n{essay}\n\n---\n\n### 🔎 Sources Used\n{refs}"
        result = f"{essay}\n\n---\n\n### 🔎 Sources Used\n{refs}"

        return result, gr.update(visible=True)
        
    except Exception as e:
        print(f"Error in generate_paper: {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        return f"Error generating paper: {str(e)}", gr.update(visible=False)

def export_paper(format_choice):
    """Export the generated paper in the selected format."""
    global last_essay, last_refs
    
    if not last_essay:
        return None
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    if format_choice == "TXT":
        filename = f"paper_{timestamp}.txt"
        filepath = os.path.join(tempfile.gettempdir(), filename)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("GENERATED ACADEMIC PAPER\n")
            f.write("=" * 50 + "\n\n")
            f.write(last_essay)
            f.write("\n\n" + "=" * 50 + "\n")
            f.write("SOURCES USED\n")
            f.write("=" * 50 + "\n")
            f.write(last_refs)
        return filepath
    
    elif format_choice == "DOCX":
        filename = f"paper_{timestamp}.docx"
        filepath = os.path.join(tempfile.gettempdir(), filename)
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Generated Academic Paper', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Essay content
        lines = last_essay.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('###'):
                doc.add_heading(line.replace('###', '').strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line.strip('*'))
                run.bold = True
            else:
                doc.add_paragraph(line)
        
        # Sources section
        doc.add_page_break()
        doc.add_heading('Sources Used', 1)
        for ref in last_refs.split('\n'):
            if ref.strip():
                doc.add_paragraph(ref.strip('- '), style='List Bullet')
        
        doc.save(filepath)
        return filepath
    
    elif format_choice == "PDF":
        filename = f"paper_{timestamp}.pdf"
        filepath = os.path.join(tempfile.gettempdir(), filename)
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Title
        pdf.set_font("Helvetica", 'B', 16)
        pdf.cell(0, 10, 'Generated Academic Paper', align='C', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(10)
        
        # Essay content
        pdf.set_font("Helvetica", '', 11)
        clean_essay = last_essay.replace('**', '').replace('###', '').replace('##', '')
        
        for line in clean_essay.split('\n'):
            if line.strip():
                pdf.multi_cell(0, 6, line.strip())
                pdf.ln(2)
        
        # Sources section
        pdf.add_page()
        pdf.set_font("Helvetica", 'B', 14)
        pdf.cell(0, 10, 'Sources Used', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(5)
        
        pdf.set_font("Helvetica", '', 10)
        for ref in last_refs.split('\n'):
            if ref.strip():
                pdf.multi_cell(0, 5, ref.strip())
                pdf.ln(2)
        
        pdf.output(filepath)
        return filepath
    
    elif format_choice == "MD":
        filename = f"paper_{timestamp}.md"
        filepath = os.path.join(tempfile.gettempdir(), filename)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("# Generated Academic Paper\n\n")
            f.write(last_essay)
            f.write("\n\n---\n\n")
            f.write("## Sources Used\n\n")
            f.write(last_refs)
        return filepath

# --------------------  GRADIO UI  --------------------
with gr.Blocks(theme="soft") as app:
    gr.Markdown("## 🧠 Academic Paper Chatbot — RAG + Hugging Face API")
    gr.Markdown("Upload your personal notes, ask a question, and generate an academic paper draft.")
    with gr.Tab("Dashboard"):
        gr.Markdown("### Pending Files (not yet in store)")
        with gr.Row():
            # -----  file drop zone  -----
            file_uploader = gr.File(file_count="multiple", label="📂 Upload files")
            with gr.Column():
                # -----  pending queue (NEW)  -----
                #gr.Markdown("### Pending Files (not yet in store)")
                pending_display = gr.Textbox(value=pending_value, label="Queued Files", interactive=False, lines=5)
                add_btn = gr.Button("📥 Add to Store", variant="secondary")

        gr.Markdown("### Files in Store")
        with gr.Row():
            with gr.Column():
                # -----  already stored  -----
                # gr.Markdown("### Files in Store")
                stored_display = gr.Textbox(value="No files in store", label="Stored Files", interactive=False, lines=8)
                status_msg = gr.Markdown()
            with gr.Column():
                # -----  remove / clear  -----
                remove_index = gr.Textbox(placeholder="index (0,1,…)", label="Remove file by index", scale=3)
                remove_btn = gr.Button("❎ Remove", variant="secondary", scale=1)
                clear_btn = gr.Button("🗑️ Clear All", variant="secondary")

        # ----------  paper generation  ----------
        gr.Markdown("### Paper Generation")
        query_input = gr.Textbox(label="🎯 Your Question / Essay Prompt", lines=3)
        gen_btn = gr.Button("🧩 Generate Paper")

        # ----------  paper download  ----------
        with gr.Row(visible=False) as export_section:
            output_md = gr.Textbox(label="📄 Generated Paper", lines=8)
            with gr.Column():
                format_dropdown = gr.Dropdown(["TXT", "DOCX", "PDF", "MD"], value="DOCX", label="Export Format")
                export_btn = gr.Button("💾 Download Paper")
                download_file = gr.File()

        # ----------  events  ----------
        file_uploader.upload(
            on_drop_more,
            inputs=[file_uploader],                     # user drop
            outputs=[pending_files, pending_display, file_uploader]  # <- CLEAR added
        ).then(lambda: None, None, file_uploader)      # extra insurance: reset value
        add_btn.click(
            process_files,
            inputs=[],                                  # we read from pending_files state
            outputs=[stored_display, status_msg, pending_display]
        )
        remove_btn.click(remove_file, inputs=[remove_index], outputs=[stored_display, status_msg])
        clear_btn.click(clear_all, inputs=[], outputs=[stored_display, status_msg])
        gen_btn.click(generate_paper, inputs=[query_input], outputs=[output_md, export_section])
        export_btn.click(export_paper, inputs=[format_dropdown], outputs=[download_file])
    
    with gr.Tab("Advanced Options"):
        gr.Markdown("### Tailor Paper Structure")
        gr.Markdown("### Load or Save Store")
        with gr.Row():
            with gr.Column():
                store_selector = gr.File(label="Select store file (*.vsto.json)")
                load_btn = gr.Button("📂 Load Store")
            with gr.Column():
                save_btn = gr.Button("💾 Save Current Store", variant="secondary")
                store_saver = gr.DownloadButton(label="⬇ Download store file", visible=False) # hidden placeholder
        gr.Markdown("### Manage LLM Models")

        # ----------  events (advanced options) ----------
        load_btn.click(load_store_file,
            inputs=[store_selector],
            outputs=[stored_display, status_msg]
            )
        save_btn.click(save_store_file,
            inputs=[],
            outputs=[store_saver]
            ).then(  # fill the hidden File component
                lambda f: gr.update(value=f, visible=True),  # make it visible for download
                inputs=[store_saver],
                outputs=[store_saver]
                )
        
if __name__ == "__main__":
    app.launch()